import os
import csv
import requests
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, MessageHandler, filters, CallbackContext, CallbackQueryHandler
from docx import Document
    
TELEGRAM_API_TOKEN = '6528605793:AAGi5qHjPY43LmG-KsahUvk2K1TmLptr6TU'
GOOGLE_BOOKS_API_KEY = 'AIzaSyAmu8yU0vuY-q1vWENu8OirNwIIDrgFh6g'

reading_list = []

def start(update: Update, context: CallbackContext):
    update.message.reply_text("Welcome to PagePal! I can help you find books. Use /help to see what I can do.")

def help_command(update: Update, context: CallbackContext):
    help_text = """
    Here are the commands you can use:
    /start - Welcome message
    /book - Get a list of books by genre
    /preview - Get a preview link for a book
    /list - Add/Delete/View your reading list
    /reading_list - Manage your reading list
    /help - Show this help message
    """
    update.message.reply_text(help_text)

def book(update: Update, context: CallbackContext):
    update.message.reply_text("Please enter the genre you are interested in:")
    context.user_data['action'] = 'genre'  

def handle_genre(update: Update, context: CallbackContext):
    genre = update.message.text
    url = f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}'
    response = requests.get(url).json()

    if 'items' in response:
        books = response['items']
        filename = f'{genre}_books.csv'
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['Title', 'Author', 'Description', 'Year Published', 'Language', 'Preview Link']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()

            for book in books:
                volume_info = book['volumeInfo']
                writer.writerow({
                    'Title': volume_info.get('title', 'N/A'),
                    'Author': ', '.join(volume_info.get('authors', ['N/A'])),
                    'Description': volume_info.get('description', 'N/A'),
                    'Year Published': volume_info.get('publishedDate', 'N/A'),
                    'Language': volume_info.get('language', 'N/A'),
                    'Preview Link': volume_info.get('previewLink', 'N/A')
                })

        update.message.reply_document(open(filename, 'rb'))
        os.remove(filename)
    else:
        update.message.reply_text("Sorry, I couldn't find books for that genre.")
    
    context.user_data.pop('action', None) 

def preview(update: Update, context: CallbackContext):
    update.message.reply_text("Please enter the name of the book you want to preview:")
    context.user_data['action'] = 'preview'  

def handle_preview(update: Update, context: CallbackContext):
    book_name = update.message.text
    url = f'https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={GOOGLE_BOOKS_API_KEY}'
    response = requests.get(url).json()

    if 'items' in response:
        book = response['items'][0]
        preview_link = book['volumeInfo'].get('previewLink', 'No preview available')
        update.message.reply_text(f'Preview link: {preview_link}')
    else:
        update.message.reply_text("Sorry, I couldn't find a preview for that book.")

    context.user_data.pop('action', None)  

def reading_list_command(update: Update, context: CallbackContext):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text('Manage your reading list:', reply_markup=reply_markup)

def handle_buttons(update: Update, context: CallbackContext):
    query = update.callback_query
    query.answer()

    if query.data == 'add':
        query.edit_message_text(text="Please enter the book title to add:")
        context.user_data['action'] = 'add'
    elif query.data == 'delete':
        query.edit_message_text(text="Please enter the book title to delete:")
        context.user_data['action'] = 'delete'
    elif query.data == 'view':
        if reading_list:
            doc = Document()
            doc.add_heading('Your Reading List', 0)

            for book in reading_list:
                doc.add_paragraph(book['Title'])
                if book['Preview Link']:
                    doc.add_paragraph(f'Preview: {book["Preview Link"]}', style='Intense Quote')

            doc_filename = 'Reading_List.docx'
            doc.save(doc_filename)
            query.message.reply_document(open(doc_filename, 'rb'))
            os.remove(doc_filename)
        else:
            query.edit_message_text(text="Your reading list is empty.")

def handle_message(update: Update, context: CallbackContext):
    text = update.message.text
    action = context.user_data.get('action')

    if action == 'genre':
        handle_genre(update, context)
    elif action == 'preview':
        handle_preview(update, context)
    elif action == 'add':
        add_to_reading_list(text, update)
    elif action == 'delete':
        remove_from_reading_list(text, update)
    else:
        update.message.reply_text("Please use one of the commands to interact with me.")

def add_to_reading_list(book_name, update):
    url = f'https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={GOOGLE_BOOKS_API_KEY}'
    response = requests.get(url).json()

    if 'items' in response:
        book = response['items'][0]
        title = book['volumeInfo'].get('title', 'N/A')
        preview_link = book['volumeInfo'].get('previewLink', 'N/A')
        reading_list.append({'Title': title, 'Preview Link': preview_link})
        update.message.reply_text(f'Added "{title}" to your reading list.')
    else:
        update.message.reply_text("Sorry, I couldn't find that book.")

def remove_from_reading_list(book_name, update):
    global reading_list
    new_list = [book for book in reading_list if book_name.lower() not in book['Title'].lower()]

    if len(new_list) < len(reading_list):
        reading_list = new_list
        update.message.reply_text(f'Removed "{book_name}" from your reading list.')
    else:
        update.message.reply_text(f'Could not find "{book_name}" in your reading list.')

def main():
    updater = Updater(TELEGRAM_API_TOKEN)
    dispatcher = updater.dispatcher

    dispatcher.add_handler(CommandHandler("start", start))
    dispatcher.add_handler(CommandHandler("help", help_command))
    dispatcher.add_handler(CommandHandler("book", book))
    dispatcher.add_handler(CommandHandler("preview", preview))
    dispatcher.add_handler(CommandHandler("reading_list", reading_list_command))
    dispatcher.add_handler(MessageHandler(Filters.text & ~Filters.command, handle_message))
    dispatcher.add_handler(CallbackQueryHandler(handle_buttons))

    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()

